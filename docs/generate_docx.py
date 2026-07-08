from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.18)
section.right_margin  = Inches(1.18)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Styles helper ────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1, center=False):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def para(doc, text="", bold=False, size=11, center=False, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, size=size, italic=italic)
    return p

def bullet(doc, text, bold_prefix=None, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        set_font(r1, bold=True, size=size)
        r2 = p.add_run(text)
        set_font(r2, size=size)
    else:
        r = p.add_run(text)
        set_font(r, size=size)
    return p

def shade_row(row, hex_color="D9E1F2"):
    """Apply background shading to all cells in a row."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

def make_table(doc, headers, rows, col_widths=None, shade_header=True):
    """Create a table with formatted header and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    if shade_header:
        shade_row(hdr, "2F5496")  # dark blue like reference
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        is_bold_row = any(
            isinstance(v, str) and v.startswith("**") for v in row_data
        )
        if is_bold_row:
            shade_row(row, "D9E1F2")
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            val_clean = str(val).replace("**", "")
            run = p.add_run(val_clean)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.bold = is_bold_row or val.startswith("**") if isinstance(val, str) else False

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
para(doc, "International School – Duy Tan University", bold=False, size=12, center=True, space_before=24)
para(doc, "Capstone Project", bold=True, size=14, center=True)
para(doc, "CMU-IS 451", bold=False, size=12, center=True)
para(doc)

heading(doc, "Project Plan", level=1, center=True)
para(doc, "Version 1.0", size=12, center=True)
para(doc, "Date: July 7th, 2026", size=12, center=True)
para(doc)

para(doc, "FoodieGo – Online Food Ordering System", bold=True, size=14, center=True)
para(doc, "Submitted by Team FoodieGo", size=12, center=True)
para(doc)

for name in ["Cao Đình Bảo", "Võ Duy Hoàng", "Phạm Hải Thiên"]:
    para(doc, name, size=12, center=True)

para(doc)
para(doc, "Approved by", bold=True, size=12, center=True)
para(doc, "Project Mentor: ______________________", size=12, center=True)
para(doc)

# Signature tables
for label in ["Project Plan Review Panel Representative:", "Capstone Project – Mentor:"]:
    para(doc, label, bold=True, size=11)
    make_table(doc, ["Name", "Signature", "Date"], [["", "", ""]], col_widths=[6, 5, 4])
    para(doc)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# PROJECT INFORMATION
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "PROJECT INFORMATION", level=1)
make_table(doc,
    ["Item", "Information"],
    [
        ["Project Acronym", "FoodieGo"],
        ["Project Title", "FoodieGo – Online Food Ordering System"],
        ["Start Date", "07 July 2026"],
        ["End Date", "30 August 2026"],
        ["Lead Institution", "International School, Duy Tan University"],
        ["Project Mentor", "__________________"],
        ["Scrum Master / Project Leader & contact details", "Cao Đình Bảo, Email: caodinhbao@dtu.edu.vn"],
        ["Partner Organization", "Duy Tan University"],
        ["Project Repository", "GitHub – FoodieGo"],
    ],
    col_widths=[7, 9]
)
para(doc)

para(doc, "Team Members:", bold=True)
make_table(doc,
    ["Name", "Student ID", "Email", "Role"],
    [
        ["Cao Đình Bảo",    "29219051113", "caodinhbao@dtu.edu.vn",    "Scrum Master"],
        ["Võ Duy Hoàng",    "29219020704", "voduyhoan@dtu.edu.vn",     "Team Member"],
        ["Phạm Hải Thiên",  "29219020597", "phamhaithien@dtu.edu.vn",  "Team Member"],
    ],
    col_widths=[4, 3.5, 5.5, 3.5]
)
para(doc)
page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# DOCUMENT NAME & REVISION HISTORY
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "DOCUMENT NAME", level=1)
make_table(doc,
    ["Item", "Information"],
    [
        ["Document Title", "Project Plan Document"],
        ["Author(s)",      "Cao Đình Bảo"],
        ["Role",           "Scrum Master"],
        ["Date",           "July 7th, 2026"],
        ["File name",      "FoodieGo_ProjectPlan_ver1.0.docx"],
    ],
    col_widths=[5, 11]
)
para(doc)

heading(doc, "REVISION HISTORY", level=1)
make_table(doc,
    ["Version", "Date", "Comments", "Author"],
    [
        ["1.0", "July 7th, 2026", "Initial Release", "Cao Đình Bảo"],
    ],
    col_widths=[2.5, 4, 6, 4]
)
para(doc)
page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "TABLE OF CONTENTS", level=1)
toc_items = [
    "1. Introduction",
    "    1.1 Purpose",
    "    1.2 Project Overview",
    "    1.3 Project Deliverable",
    "2. Team Organization",
    "    2.1 Scrum Team Information",
    "    2.2 Role and Responsibility",
    "    2.3 Communication Methodology",
    "    2.4 Communication and Report",
    "3. Development Process",
    "4. Schedule and Cost",
    "    4.1 Detailed Schedule",
    "    4.2 Cost",
    "5. Project Risk",
    "6. Deliverables",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

para(doc)
heading(doc, "LIST OF TABLES", level=1)
lot_items = [
    "Table 1. Scrum Team Organization",
    "Table 2. Role and Responsibilities",
    "Table 3. Communication Methodology",
    "Table 4. Communication and Report",
    "Table 5. Detailed Schedule",
    "Table 6. Cost person/hours",
    "Table 7. Total Cost Estimate (1)",
    "Table 8. Cost Estimation Parameters",
    "Table 9. Total Working Hours",
    "Table 10. Labor Cost Distribution",
    "Table 11. Other Project Costs",
    "Table 12. Total Estimated Cost Of The Project",
    "Table 13. Rating for likelihood and seriousness for each risk",
    "Table 14. Project Risk",
    "Table 15. Deliverables",
]
for item in lot_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Introduction", level=1)

heading(doc, "1.1 Purpose", level=2)
para(doc, (
    "This document provides a summary of the project objectives, scope, and work assignments. "
    "It outlines key milestones, required resources, an overall schedule, and a budget plan. "
    "The Project Plan is developed based on the approved proposal of FoodieGo – Online Food Ordering "
    "System to ensure that the system is implemented on schedule, meets stakeholder requirements, "
    "and aligns with the planned deliverables."
), size=11)

heading(doc, "1.2 Project Overview", level=2)
para(doc, (
    "FoodieGo is an online food ordering platform that enables customers to browse restaurants, place "
    "orders, track deliveries, and review restaurants. Restaurant owners can manage restaurants, menus, "
    "and orders. Administrators monitor users, restaurants, and system statistics. The project follows "
    "Agile Scrum methodology with Lean principles to reduce waste and shorten Lead Time."
), size=11)

heading(doc, "1.3 Project Deliverable", level=2)
para(doc, (
    "The project will develop a web-based food ordering system designed for three types of users: "
    "customers, restaurant owners, and administrators. The system will support core functions such as "
    "user authentication, restaurant browsing, menu management, order placement, real-time delivery fee "
    "calculation (via FastAPI microservice), and order tracking. Quality assurance will be maintained "
    "through CI/CD pipelines (GitHub Actions), automated testing (Jest), and code quality analysis (SonarQube)."
), size=11)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 2. TEAM ORGANIZATION
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Team Organization", level=1)

heading(doc, "2.1 Scrum Team Information", level=2)
para(doc, "Table 1. Scrum Team Organization", bold=True, size=10, center=True)
make_table(doc,
    ["Full Name", "Email", "Position"],
    [
        ["__________________",  "__________________",          "Mentor"],
        ["Cao Đình Bảo",        "caodinhbao@dtu.edu.vn",       "Scrum Master"],
        ["Võ Duy Hoàng",        "voduyhoan@dtu.edu.vn",        "Team Member"],
        ["Phạm Hải Thiên",      "phamhaithien@dtu.edu.vn",     "Team Member"],
    ],
    col_widths=[5, 7, 4]
)
para(doc)

heading(doc, "2.2 Role and Responsibility", level=2)
para(doc, "Table 2. Role and Responsibilities", bold=True, size=10, center=True)
make_table(doc,
    ["Role", "Responsibility", "Name/Title"],
    [
        ["Scrum Master",  "Organize Scrum events (Planning, Daily, Review, Retro). Ensure team compliance with Scrum. Remove impediments, support coordination.", "Cao Đình Bảo"],
        ["Secretary",     "Record meeting minutes. Archive documents (Proposal, Project Plan, Report). Track meeting schedules and deadlines.", "Cao Đình Bảo"],
        ["Reviewer",      "Participate in product evaluation in Sprint Review. Provide feedback from user perspective. Ensure features meet needs.", "All Members"],
        ["Developer",     "Design and program frontend interface. Build API and database (Backend). Integrate microservices.", "All Members"],
        ["Analyzer",      "Analyze user requirements. Identify use cases, user stories. Support Scrum Master in writing backlog.", "All Members"],
        ["Tester",        "Write Test Plan and Test Case. Test software (unit test, integration test). Record and report bugs.", "All Members"],
        ["Mentor",        "Project direction and monitoring. Provide professional advice. Support in solving technical difficulties.", "Project Mentor"],
    ],
    col_widths=[3.5, 9, 3.5]
)
para(doc)

heading(doc, "2.3 Communication Methodology", level=2)
para(doc, "Table 3. Communication Methodology", bold=True, size=10, center=True)
make_table(doc,
    ["Audience / Attendees", "Topic / Deliverable", "Frequency", "Method"],
    [
        ["Mentor and Team Member", "Project Progress Review",                    "Weekly", "Meeting, Email, Zalo"],
        ["Team Member",            "Project Progress Review and Daily Meeting",   "Daily",  "GitHub, Discord, Zalo, Google Meet"],
    ],
    col_widths=[4.5, 5.5, 2.5, 4]
)
para(doc)

heading(doc, "2.4 Communication and Report", level=2)
para(doc, "Table 4. Communication and Report", bold=True, size=10, center=True)
make_table(doc,
    ["Type of communication", "Methods, tools", "Frequency", "Information", "People"],
    [
        ["Scrum meeting (Daily Stand-up)", "Discord / Face to face", "Every day",
         "What was done in the last 24 hours, working on plans for today, difficulties and solutions. 10–15 minutes.",
         "Project team"],
        ["Sprint Planning Meeting", "Google Meet / Face to face", "Every 2 weeks",
         "All members analyze requirements, plan and design the sprint going to do.",
         "Project team"],
        ["Retrospective Meeting",   "Google Meet / Face to face", "Every 2 weeks",
         "Sharing strengths and weaknesses. Period for each member and the solution measured.",
         "Project team & Mentor"],
        ["Demo",                    "Google Meet (online)",       "Every Sprint",
         "Demonstrate working software to mentor and stakeholders. Collect feedback.",
         "Project team & Mentor"],
        ["Task Tracking",           "GitHub Projects",            "Daily",
         "Web-based task tracking system. Manage / divide tasks, report bugs/issues.",
         "Project team"],
    ],
    col_widths=[3.5, 3, 2.5, 5.5, 2.5]
)
para(doc)
page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 3. DEVELOPMENT PROCESS
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Development Process", level=1)
para(doc, (
    "Scrum is an agile project management framework that helps teams structure and manage their work "
    "through a set of values, principles, and practices. FoodieGo applies Scrum combined with Lean "
    "principles to reduce waste and shorten Lead Time."
), size=11)

para(doc, "Benefits of Scrum applied in FoodieGo:", bold=True)
scrum_benefits = [
    ("Adapts Quickly to Change:", "Scrum allows the team to flexibly adjust plans and goals during short work cycles (Sprints). This helps the project easily react to changes without causing major disruptions."),
    ("Identifies Problems Early:", "Daily Stand-ups and end-of-Sprint reviews provide opportunities for the team to detect and resolve problems early, reducing bug-fixing costs and preventing delays."),
    ("Focuses on Value:", "Scrum prioritizes developing the features with the highest value first. The team doesn't waste time on less important requirements."),
    ("Better Aligns with Customer Needs:", "With frequent Sprint Reviews, the team can receive direct feedback and ensure the product is always adjusted to meet user expectations."),
    ("Increases Work Productivity:", "Scrum fosters a transparent and collaborative work environment. Team members know their tasks, share information continuously, and work more efficiently."),
    ("Predictable Scheduling:", "Fixed-length Sprints maintain a stable work rhythm. Clients and stakeholders can easily predict progress and plan feature releases."),
]
for bold_part, rest in scrum_benefits:
    bullet(doc, rest, bold_prefix=bold_part)

para(doc, "Lean principles applied:", bold=True)
lean_items = [
    "Reduce waiting time between tasks",
    "Reduce unnecessary rework through PR reviews and CI/CD",
    "Measure Lead Time from code commit to merge",
    "Continuous Improvement via Sprint Retrospective",
]
for item in lean_items:
    bullet(doc, item)

para(doc, "Quality is measured by:", bold=True)
quality_items = [
    "Test Coverage ≥ 80% (Jest)",
    "CI Success Rate ≥ 90% (GitHub Actions)",
    "SonarQube Issues ≤ 10",
    "Lead Time (PR open → merge) ≤ 1 day",
]
for item in quality_items:
    bullet(doc, item)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 4. SCHEDULE AND COST
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Schedule and Cost", level=1)
heading(doc, "4.1 Detailed Schedule", level=2)
para(doc, "Table 5. Detailed Schedule", bold=True, size=10, center=True)

schedule_rows = [
    ["**1", "**Project Initiation",                   "**07/07", "**11/07", "**120"],
    ["1.1",  "Gathering Requirement",                  "07/07",   "08/07",   "40"],
    ["1.1.1","Get requirement from Mentor",            "07/07",   "07/07",   "15"],
    ["1.1.2","Analyzing requirement",                  "08/07",   "08/07",   "25"],
    ["1.2",  "Create Proposal Document",               "09/07",   "11/07",   "80"],
    ["1.2.1","Product Definition",                     "09/07",   "09/07",   "20"],
    ["1.2.2","Business Need",                          "09/07",   "10/07",   "20"],
    ["1.2.3","Proposed Solution",                      "10/07",   "11/07",   "25"],
    ["1.2.4","Master Plan",                            "11/07",   "11/07",   "15"],
    ["**2", "**Project Setup",                        "**12/07", "**16/07", "**120"],
    ["2.1.1","Project Kick-off Meeting",               "12/07",   "12/07",   "10"],
    ["2.1.2","Team Meeting – Clarify Scope",           "12/07",   "13/07",   "15"],
    ["2.1.3","Create User Stories",                    "13/07",   "14/07",   "30"],
    ["2.1.4","Create Product Backlog",                 "14/07",   "15/07",   "30"],
    ["2.1.5","Review Document",                        "15/07",   "16/07",   "20"],
    ["2.1.6","Create Project Plan",                    "16/07",   "16/07",   "15"],
    ["**3", "**Development",                          "**17/07", "**28/08", "**972"],
    ["**3.1","**Sprint 1 (Auth + Users)",             "**17/07", "**30/07", "**252"],
    ["3.1.1","Sprint Planning",                        "17/07",   "17/07",   "15"],
    ["3.1.2","Setup Infrastructure & DB",              "18/07",   "20/07",   "54"],
    ["3.1.3","Authentication & JWT",                   "21/07",   "25/07",   "90"],
    ["3.1.4","Profile & Role Management",              "26/07",   "29/07",   "72"],
    ["3.1.5","Sprint Release",                         "30/07",   "30/07",   "21"],
    ["**3.2","**Sprint 2 (Restaurant + Menu)",        "**31/07", "**13/08", "**252"],
    ["3.2.1","Sprint Planning",                        "31/07",   "31/07",   "15"],
    ["3.2.2","Restaurant Profile APIs",                "01/08",   "05/08",   "72"],
    ["3.2.3","Menu Items CRUD",                        "06/08",   "10/08",   "90"],
    ["3.2.4","Database Unit Testing",                  "11/08",   "13/08",   "54"],
    ["3.2.5","Sprint Release",                         "13/08",   "13/08",   "21"],
    ["**3.3","**Sprint 3 (Orders + Delivery)",        "**14/08", "**23/08", "**252"],
    ["3.3.1","Sprint Planning",                        "14/08",   "14/08",   "15"],
    ["3.3.2","Order Creation & Status APIs",           "15/08",   "18/08",   "90"],
    ["3.3.3","Delivery Service (FastAPI)",             "19/08",   "21/08",   "90"],
    ["3.3.4","Integration Testing",                    "22/08",   "23/08",   "36"],
    ["3.3.5","Sprint Release",                         "23/08",   "23/08",   "21"],
    ["**3.4","**Sprint 4 (Frontend + CI/CD + Docs)", "**24/08", "**28/08", "**216"],
    ["3.4.1","Sprint Planning",                        "24/08",   "24/08",   "15"],
    ["3.4.2","Web Interface Integration",              "25/08",   "26/08",   "108"],
    ["3.4.3","Docker & CI/CD Pipeline",               "27/08",   "27/08",   "54"],
    ["3.4.4","Final Release & Demo",                  "28/08",   "28/08",   "39"],
    ["**4", "**Close Project",                        "**29/08", "**30/08", "**36"],
    ["4.1",  "Final Release",                          "29/08",   "29/08",   "12"],
    ["4.2",  "Project Meeting",                        "29/08",   "30/08",   "12"],
    ["4.3",  "Final Submission",                       "30/08",   "30/08",   "12"],
]

table = doc.add_table(rows=1 + len(schedule_rows), cols=5)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0]
shade_row(hdr, "2F5496")
for i, h in enumerate(["No.", "Task Name", "Start", "Finish", "Effort (Hours)"]):
    p = hdr.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.font.name = "Times New Roman"; run.font.size = Pt(10)
    run.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)

for ri, row_data in enumerate(schedule_rows):
    row = table.rows[ri + 1]
    is_bold = row_data[0].startswith("**")
    if is_bold:
        shade_row(row, "D9E1F2")
    for ci, val in enumerate(row_data):
        cell = row.cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        clean = val.replace("**", "")
        run = p.add_run(clean)
        run.font.name = "Times New Roman"; run.font.size = Pt(10)
        run.bold = is_bold

# set col widths
col_widths_schedule = [Cm(1.5), Cm(8), Cm(2), Cm(2), Cm(3)]
for i, w in enumerate(col_widths_schedule):
    for row in table.rows:
        row.cells[i].width = w

para(doc)
page_break(doc)

# ── COST ────────────────────────────────────────────────────────────────────
heading(doc, "4.2 Cost", level=2)
para(doc, (
    "This section presents the estimated cost required for the development of the FoodieGo project. "
    "The total cost includes labor costs (based on working hours of team members) and additional "
    "operational costs such as hosting and infrastructure."
), size=11)

heading(doc, "4.2.1 Cost per Person/Hour", level=3)
para(doc, (
    "The labor cost of the project is estimated based on the hourly rate of each team member. "
    "The proposed hourly rate is aligned with the typical internship salary in Vietnam, which ranges "
    "from 70,000 VND to 100,000 VND per hour. The Scrum Master is assigned a slightly higher hourly "
    "rate due to additional responsibilities such as project coordination, sprint planning, and team "
    "management, in addition to development tasks."
), size=11)

para(doc, "Table 6. Cost person/hours", bold=True, size=10, center=True)
make_table(doc,
    ["Full Name", "Role", "Salary Rate (USD/hour)"],
    [
        ["Võ Duy Hoàng",   "Team Member",  "3.50"],
        ["Phạm Hải Thiên", "Team Member",  "3.50"],
        ["Cao Đình Bảo",   "Scrum Master", "4.00"],
    ],
    col_widths=[6, 5, 5.5]
)
para(doc)

heading(doc, "4.2.2 Total Cost Estimate", level=3)
para(doc, "The total project cost consists of two main components:")
bullet(doc, "Labor cost based on the working hours of team members")
bullet(doc, "Other operational costs required to support project development and deployment")

para(doc, "Table 7. Total Cost Estimate (1)", bold=True, size=10, center=True)
make_table(doc,
    ["No.", "Criteria", "Price (USD)"],
    [
        ["1", "Working hours", "4,576.00"],
        ["2", "Other cost",    "190.00"],
        ["**TOTAL", "**ESTIMATED COST", "**4,766.00"],
    ],
    col_widths=[2, 10, 4.5]
)
para(doc)

heading(doc, "4.2.3 Project Cost Parameters", level=3)
para(doc, "The parameters used to estimate the total project cost are presented in Table 8.", size=11)
para(doc, "Table 8. Cost Estimation Parameters", bold=True, size=10, center=True)
make_table(doc,
    ["Description", "Amount", "Unit"],
    [
        ["Number of members",                  "3",    "Person"],
        ["Number of working hours per day",    "6",    "Hours"],
        ["The average cost per hour per member","3.67", "USD"],
        ["The number of working days",         "54",   "Days"],
    ],
    col_widths=[9, 3, 4.5]
)
para(doc)

heading(doc, "4.2.4 Total Working Hours", level=3)
para(doc, "The total estimated working hours for the project are 1,248 hours, derived from the detailed activities defined in the Project Plan.", size=11)
para(doc, "Table 9. Total Working Hours", bold=True, size=10, center=True)
make_table(doc,
    ["Phase", "Working Hours"],
    [
        ["Project Initiation",                  "120 hours"],
        ["Project Setup",                       "120 hours"],
        ["Sprint 1 (Auth + Users)",             "252 hours"],
        ["Sprint 2 (Restaurant + Menu)",        "252 hours"],
        ["Sprint 3 (Orders + Delivery)",        "252 hours"],
        ["Sprint 4 (Frontend + CI/CD + Docs)",  "216 hours"],
        ["Close Project",                       "36 hours"],
        ["**TOTAL",                             "**1,248 hours"],
    ],
    col_widths=[10, 6.5]
)
para(doc)

heading(doc, "4.2.5 Labor Cost Calculation", level=3)
para(doc, "The labor cost is calculated based on the actual working hours of each team member multiplied by their hourly rate. Since all three members participate in the development process, the workload is distributed equally.", size=11)
para(doc, "Table 10. Labor Cost Distribution", bold=True, size=10, center=True)
make_table(doc,
    ["Full Name", "Role", "Salary Rate (USD/hour)", "Hours", "Amount (USD)"],
    [
        ["Võ Duy Hoàng",   "Team Member",  "3.50", "416", "1,456.00"],
        ["Phạm Hải Thiên", "Team Member",  "3.50", "416", "1,456.00"],
        ["Cao Đình Bảo",   "Scrum Master", "4.00", "416", "1,664.00"],
        ["**TOTAL",        "",             "",     "**1,248", "**4,576.00"],
    ],
    col_widths=[4, 3.5, 4, 2.5, 3]
)
para(doc)

heading(doc, "4.2.6 Other Costs", level=3)
para(doc, "Besides labor costs, the project requires several supporting operational expenses, including infrastructure and miscellaneous costs.", size=11)
para(doc, "Table 11. Other Project Costs", bold=True, size=10, center=True)
make_table(doc,
    ["Item", "Cost (USD)", "Note"],
    [
        ["Electricity and internet", "70",   "Team working environment during the project"],
        ["Server / hosting rental",  "100",  "Deployment and testing environment"],
        ["Domain (optional)",        "20",   "FoodieGo domain registration"],
        ["**TOTAL",                  "**190",""],
    ],
    col_widths=[5, 3, 8.5]
)
para(doc)

heading(doc, "4.2.7 Total Project Cost", level=3)
para(doc, "The total estimated cost of the project is calculated as follows:", size=11)
para(doc, "Table 12. Total Estimated Cost Of The Project", bold=True, size=10, center=True)
make_table(doc,
    ["Item", "Amount (USD)"],
    [
        ["Labor cost",             "4,576.00"],
        ["Other costs",            "190.00"],
        ["**TOTAL ESTIMATED COST", "**4,766.00"],
    ],
    col_widths=[10, 6.5]
)
para(doc)
p = doc.add_paragraph()
r1 = p.add_run("Therefore, the total estimated budget required to complete the FoodieGo project is ")
r1.font.name = "Times New Roman"; r1.font.size = Pt(11)
r2 = p.add_run("4,766.00 USD")
set_font(r2, bold=True, size=11)
r3 = p.add_run(" (In words: Four thousand seven hundred sixty-six US dollars).")
r3.font.name = "Times New Roman"; r3.font.size = Pt(11)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 5. PROJECT RISK
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Project Risk", level=1)

para(doc, "Table 13. Rating for likelihood and seriousness for each risk", bold=True, size=10, center=True)
make_table(doc,
    ["RATING FOR LIKELIHOOD AND SERIOUSNESS FOR EACH RISK", ""],
    [
        ["L",  "Rated as Low"],
        ["E",  "Rated as Extreme (Used for Seriousness only)"],
        ["M",  "Rated as Medium"],
        ["NA", "Not Assessed"],
        ["H",  "Rated as High"],
    ],
    col_widths=[3, 13.5]
)
para(doc)

para(doc, "Table 14. Project Risk", bold=True, size=10, center=True)
make_table(doc,
    ["Risk", "Definition", "Level", "Likelihood", "Mitigation Strategy"],
    [
        ["Requirements Definition",       "The team does not have a clear agreement on functional requirements, leading to scope misalignment.", "H", "M", "Conduct requirement clarification meetings; document and confirm requirements officially before each Sprint."],
        ["Schedule Estimation",           "Planning and execution time are not accurately estimated, leading to delays.", "M", "M", "Create a detailed plan with effort estimates, update the schedule regularly via GitHub Projects."],
        ["Programming Experience",        "Team members are not familiar with technologies such as Node.js, FastAPI, or PostgreSQL.", "M", "M", "Organize internal training; pair less experienced members with experienced ones; study documentation early."],
        ["Database Management",           "Errors in PostgreSQL design could lead to data loss or low performance.", "H", "M", "Design a clear schema, perform regular backups, and test thoroughly via Docker."],
        ["Testing and Quality Assurance", "Lack of sufficient testing could cause bugs in production.", "H", "M", "Write Jest unit/integration tests with coverage ≥ 80%; integrate SonarQube for code quality checks."],
        ["Network / Hosting",             "Unstable hosting or limited bandwidth could slow down or interrupt the system.", "M", "L", "Choose a reliable hosting service; use Docker for consistent deployment environments."],
        ["Security",                      "Vulnerabilities such as SQL Injection or weak JWT handling due to insecure coding.", "H", "M", "Follow secure coding standards; validate all inputs; use HTTPS; enforce JWT expiration policies."],
        ["Team Communication",            "Lack of clear communication leads to misunderstandings or duplicated tasks.", "M", "L", "Hold daily stand-ups on Discord; use GitHub Projects for task tracking."],
        ["CI/CD Failure",                 "GitHub Actions pipeline failures block deployment and slow down development.", "M", "M", "Maintain clean CI configuration; fix lint/test errors immediately; test Docker deployment each Sprint."],
        ["Scope Creep",                   "New requirements are frequently added during development, affecting progress.", "H", "M", "Define a clear scope at the start and apply a change approval process for additional requests."],
    ],
    col_widths=[3.5, 5, 1.5, 2, 4.5]
)
para(doc)
page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# 6. DELIVERABLES
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Deliverables", level=1)
para(doc, "Table 15. Deliverables", bold=True, size=10, center=True)
make_table(doc,
    ["No.", "Document", "Deadline", "File Name"],
    [
        ["1",  "Project Proposal",          "July 10th, 2026",    "FoodieGo_ProjectProposal_ver1.0.docx"],
        ["2",  "Product Backlog",           "July 15th, 2026",    "FoodieGo_Backlog.md"],
        ["3",  "Project Plan Document",     "July 20th, 2026",    "FoodieGo_ProjectPlan_ver1.0.docx"],
        ["4",  "Sprint 1 Release",          "July 30th, 2026",    "GitHub Release – Sprint 1"],
        ["5",  "Sprint 2 Release",          "August 13th, 2026",  "GitHub Release – Sprint 2"],
        ["6",  "Sprint 3 Release",          "August 23rd, 2026",  "GitHub Release – Sprint 3"],
        ["7",  "Testing & Metrics Report",  "August 28th, 2026",  "FoodieGo_Metrics_Report.md"],
        ["8",  "Source Code (Final)",       "August 30th, 2026",  "GitHub Repository – FoodieGo"],
        ["9",  "Presentation Slides",       "August 30th, 2026",  "FoodieGo_FinalPresentation.pptx"],
        ["10", "Final Demonstration",       "August 30th, 2026",  "Live Demo – Capstone Day"],
    ],
    col_widths=[1.5, 4.5, 4, 6.5]
)

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
output_path = r"C:\Users\DINH TAO\.gemini\antigravity-ide\scratch\foodiego\docs\FoodieGo_ProjectPlan_ver1.0.docx"
doc.save(output_path)
print(f"SUCCESS: {output_path}")
